''' word from python '''
from docx import Document
from docx.shared import Cm
from docx.shared import Inches, Mm
import quotation
from datetime import datetime
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn 
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pylovepdf.ilovepdf import ILovePdf
import shutil

import os

import subprocess

def _set_cell_text_and_alignment(cell, text, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment

def _set_vertical_alignment(cell, alignment):
    try:
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        val = OxmlElement('w:vAlign')
        val.set(qn('w:val'), alignment)  # 'center', 'top', or 'bottom'
        tcPr.append(val)
    except Exception as e:
        print(e)


def _setColumn(document: Document, category: str):
    document.add_paragraph(category)
    # Assuming _set_cell_text_and_alignment and _set_vertical_alignment are previously defined

    table = document.add_table(rows=1, cols=10)
    column_widths = [Cm(1), Cm(15), Cm(1), Cm(1), Cm(1), Cm(1), Cm(15), Cm(1), Cm(1), Cm(1)]
    cell_texts = ['重量', '產品', '單價', '單位', '倉位', '重量', '產品', '單價', '單位', '倉位']
    
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = column_widths[idx]
            _set_cell_text_and_alignment(cell, cell_texts[idx])  # Assuming this is a custom function you've defined
            _set_vertical_alignment(cell, 'center')  # Assuming this is a custom function you've defined
            
            # Access the cell properties
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            
            # Define and apply borders to the cell
            for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border_element = OxmlElement(f'w:{border}')
                border_element.set(qn('w:val'), 'single')  # Apply a single line border
                border_element.set(qn('w:sz'), '4')  # Size of border
                border_element.set(qn('w:space'), '0')  # No spacing
                border_element.set(qn('w:color'), '000000')  # Black color
                tcPr.append(border_element)
    
    
    return table

def _get_unique_categories(df):
    """
    Returns a list of unique categories from the 'category' column in the DataFrame.

    Parameters:
    - df: pandas.DataFrame containing a 'category' column.

    Returns:
    - List of unique categories.
    """
    # Ensure 'category' column exists
    if 'category' in df.columns:
        unique_categories = df['category'].unique()
        return list(unique_categories)
    else:
        raise ValueError("DataFrame does not contain a 'category' column.")

def _set_cell_borders(cell):
    """Apply border settings to a table cell."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    # Define and apply borders
    for border in ['top', 'left', 'bottom', 'right']:
        border_element = OxmlElement(f'w:{border}')
        border_element.set(qn('w:val'), 'single')
        border_element.set(qn('w:sz'), '4')  # Border size
        border_element.set(qn('w:space'), '0')
        border_element.set(qn('w:color'), '000000')  # Black color
        tcPr.append(border_element)

def _filter_df(df, category):

    # Filter the DataFrame by category
    filtered_df = df[df['category'] == category]
    return filtered_df

def _set_paragraph_spacing_to_zero(paragraph):
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1

def _apply_top_bottom_borders_to_cell(cell):
    """Apply only top and bottom borders to a specific table cell."""
    # Access the cell's XML element and its properties
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    # Apply top and bottom borders
    for border in ['top', 'bottom']:
        border_element = OxmlElement(f'w:{border}')
        border_element.set(qn('w:val'), 'single')  # Border style
        border_element.set(qn('w:sz'), '4')  # Border size
        border_element.set(qn('w:space'), '0')
        border_element.set(qn('w:color'), '000000')  # Border color
        tcPr.append(border_element)

    # Remove left, right, and any other borders by setting them to 'nil'
    for border in ['left', 'right', 'insideH', 'insideV']:
        border_element = OxmlElement(f'w:{border}')
        border_element.set(qn('w:val'), 'nil')
        tcPr.append(border_element)

def _set_cell_text(cell, text, font_size=9):
    """Set the text for a cell with a specific font size."""
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(font_size)

def _remove_left_right_margins_from_cells(table):
    for row in table.rows:
        for cell in row.cells:
            # Access the cell properties
            tcPr = cell._tc.get_or_add_tcPr()
            
            # Create a new TableCellMargin element
            tcMar = OxmlElement('w:tcMar')
            
            # Set the left margin to 0
            left_margin = OxmlElement('w:left')
            left_margin.set(qn('w:w'), "0")
            left_margin.set(qn('w:type'), 'dxa')  # 'dxa' is twentieths of a point
            tcMar.append(left_margin)
            
            # Set the right margin to 0
            right_margin = OxmlElement('w:right')
            right_margin.set(qn('w:w'), "0")
            right_margin.set(qn('w:type'), 'dxa')
            tcMar.append(right_margin)
            
            # Append the modified margins back to the cell properties
            tcPr.append(tcMar)

def _split_text_evenly(text):
    """
    Splits the text into two parts at the nearest space to the midpoint.
    Returns a tuple containing the two parts.
    """
    midpoint = len(text) // 2  # Find the approximate midpoint
    
    # Search for the nearest space to the midpoint
    left = text.rfind(' ', 0, midpoint)
    right = text.find(' ', midpoint)
    
    # Determine the closest split point; prefer left if equidistant
    if right < 0 or (left >= 0 and midpoint - left <= right - midpoint):
        split_point = left
    else:
        split_point = right
    
    if split_point >= 0:
        # Split the text into two parts
        return text[:split_point], text[split_point+1:]
    else:
        # No space found, return the text as is in the first part
        return text, ''

def _table_add_row(table):
    new_row = table.add_row()
    for cell in new_row.cells:
        _set_cell_borders(cell)


def update_document_with_products(document, df, categoryList):
    MAX_ROWS_PER_PAGE = 55  # Max rows on one side before switching to the other side
    
    for category in categoryList:
        table = _setColumn(document, category )
        
        document.add_page_break() 
        
        filtered_df = _filter_df(df, category)
        
        # Track total rows filled to decide on adding a new page
        current_row_count = 0
        left_row_count = 0
        right_row_count = 0
        left_side = True
        previous_product_tag = None

        for index, row in filtered_df.iterrows():
            current_product_tag = row['productTag']


            if left_row_count < MAX_ROWS_PER_PAGE:
                cell_index_base = 0 #left side
                _table_add_row(table)
                
                left_side = True
                left_row_count += 1
                if previous_product_tag is not None and current_product_tag != previous_product_tag:
                    #for left side add another new row
                    _table_add_row(table)
                    left_row_count += 1


            elif left_row_count >= MAX_ROWS_PER_PAGE:
                cell_index_base = 5
                left_side =  False
                right_row_count += 1
                if previous_product_tag is not None and current_product_tag != previous_product_tag:
                    #for right row count + 1
                    right_row_count += 1

                
                # Note: No new row is added here; we reuse rows for the right side
            
            current_row_count = left_row_count if left_side else right_row_count

            if left_side or (left_side == False and right_row_count < MAX_ROWS_PER_PAGE):
                # Assign product details to cells
                
                table.rows[current_row_count].cells[cell_index_base].text = str(row['packing']) if row['packing'] is not None else ""
                
                concatenated_text = ''
                word_count = 0 
                for col in ['origin', 'brand', 'productTag', 'spec1', 'spec2']:
                    cell_text = str(row[col]) if row[col] is not None else ""
                    concatenated_text += cell_text + " " if col != 'packing' else cell_text
                    word_count += len(cell_text)
                    
                if word_count > 20 and left_side:
                    part1, part2 = _split_text_evenly(concatenated_text)
                    table.rows[current_row_count].cells[cell_index_base + 1].text = part1
                    _table_add_row(table)
                    table.rows[current_row_count + 1].cells[cell_index_base + 1].text = part2
                    left_row_count += 1
                elif(word_count > 20 and left_side != True):
                    part1, part2 = _split_text_evenly(concatenated_text)
                    table.rows[current_row_count].cells[cell_index_base + 1].text = part1
                    table.rows[current_row_count + 1].cells[cell_index_base + 1].text = part2
                    right_row_count += 1
                    
                    
                else:
                    table.rows[current_row_count].cells[cell_index_base + 1].text = concatenated_text

                
                # Assign other product details to subsequent cells
                table.rows[current_row_count].cells[cell_index_base + 2].text = str(row['price']) if row['price'] is not None else ""
                table.rows[current_row_count].cells[cell_index_base + 3].text = str(row['weightUnit']) if row['weightUnit'] is not None else ""
                table.rows[current_row_count].cells[cell_index_base + 4].text = str(row['warehouse']) if row['warehouse'] is not None else ""


                
                previous_product_tag = current_product_tag
                current_row_count += 1


            # Add a page break and reset counters if the current side is fully filled
            if left_row_count >= MAX_ROWS_PER_PAGE and right_row_count >= MAX_ROWS_PER_PAGE:
                 
                table =_setColumn(document,category)
                left_row_count = 0
                right_row_count = 0
                left_side = True

        _remove_left_right_margins_from_cells(table)
        

def _convert_docx_to_pdf_pandoc(input_path, output_path):
    try:
        # Specify the full pandoc command with options
        cmd = [
            'pandoc', input_path,
            '-o', output_path,
            '--pdf-engine=xelatex',
            '--template=mytemplate.latex',
            '-V', 'mainfont=Noto Serif CJK TC',
            '-V', 'documentclass=ctexart'
        ]
        
        # Execute the pandoc command
        result = subprocess.run(cmd, check=True, capture_output=True, text=True)
        
        # If the command was successful, print this message
        print(f"Conversion successful: {output_path}")
        
        # Optionally, print stdout and stderr for debugging
        print("STDOUT:", result.stdout)
        print("STDERR:", result.stderr)
        
    except subprocess.CalledProcessError as e:
        # If an error occurred during conversion, print the error message
        print(f"Error during conversion: {e}")
        print("STDOUT:", e.stdout)
        print("STDERR:", e.stderr)

def _convert_to_pdf(input_file, output_dir):
    command = [
        "soffice",
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        output_dir,
        input_file
    ]
    try:
        subprocess.run(command, check=True)
        print("Conversion successful!")
    except subprocess.CalledProcessError as e:
        print("Conversion failed:", e)

def _convert_and_rename_docx_to_pdf(api_key, input_path, desired_output_path):
    try:
        ilovepdf = ILovePdf(public_key=api_key, verify_ssl=True)
        task = ilovepdf.new_task('officepdf')
        task.add_file(input_path)
        task.execute()

        # Download the file. The download method now returns the name of the downloaded file.
        downloaded_file_name = task.download()
        
        # Assuming the downloaded file name is just the name, not the path,
        # and it's downloaded to the current working directory.
        downloaded_file_path = os.path.join(os.getcwd(), downloaded_file_name)
        
        # Move the downloaded file to the desired output path
        if os.path.exists(desired_output_path):
            os.remove(desired_output_path)  # Remove if the target file already exists
        shutil.move(downloaded_file_path, desired_output_path)

        print(f"Conversion successful. PDF saved as: {desired_output_path}")
        
    except Exception as e:
        print(f"Error during conversion: {e}")

def createQuotation(connection,effectiveDate:datetime,days: int = 2) -> str :
    document = Document()

    for paragraph in document.paragraphs:
    # Set paragraph spacing to single
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = Pt(0)

    # Set font
    style = document.styles['Normal']
    style.font.name = 'GungSeo'
    style.font.size = Pt(9)
    style.font.color.rgb = RGBColor(0, 0, 0) 
    section = document.sections[0]  # Assuming you're changing the first section
    section.page_height = Mm(297)
    section.page_width = Mm(210)

    # Optional: Set margins if you want
    section.top_margin = Inches(0.01)
    section.bottom_margin = Inches(0.25)
    section.left_margin = Inches(0.25)
    section.right_margin = Inches(0.25)

    # Add header
    header = document.sections[0].header

    # Clear existing paragraphs
    for section in document.sections:
    # Set header distance to 0
        section.header_distance = Pt(0)
        section.top_margin = Pt(0)
        # Set footer distance to 0
        section.footer_distance = Pt(0)

    for paragraph in header.paragraphs:
        paragraph.clear()

    # Company name, centered
    company_name_paragraph = header.add_paragraph('新樂食品貿易有限公司')
    company_name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in company_name_paragraph.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)  # Black font

    # Prepare the header text with placeholders for tabs
    today_date = datetime.today().strftime('%Y-%m-%d')  # Format today's date
    header_text = (
        '上水龍豐花園30號地舖|tradeasychain@gmail.com| [落單]張小姐 6045 7604/曾先生 5977 9085\n'
        '*貨品價格如有更改,恕不另行通告,價格為入倉提貨價,如有疑問請跟營業員聯絡\n'
        '*本公司只提供 <其士倉> 提貨送貨服務 5件起送 HKD$20/件\n'
        '*本公司暫不設加工服務\n'
        '*請提前落<隔夜單>以免提貨出現問題 截單時間為3:00pm\t打印日期：{}'.format(today_date)
    )

    # Add header text and configure tab stop for print date
    header_paragraph = header.add_paragraph(header_text)
    tab_stops = header_paragraph.paragraph_format.tab_stops
    tab_stop_position = Inches(6.5)  # Adjust based on your document's layout
    tab_stop = tab_stops.add_tab_stop(tab_stop_position, alignment=WD_TAB_ALIGNMENT.RIGHT)

    # Set all text to black
    for paragraph in header.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)


    df = quotation.getBestQuote(connection,effectiveDate,days)
    df.fillna('', inplace=True)
    
    categoryList = _get_unique_categories(df)
    update_document_with_products(document,df,categoryList)
    #single spacing
    for paragraph in document.paragraphs:
        _set_paragraph_spacing_to_zero(paragraph)

    # Iterate through headers and footers in all sections
    for section in document.sections:
        for header in section.header.paragraphs:
            _set_paragraph_spacing_to_zero(header)
        for footer in section.footer.paragraphs:
            _set_paragraph_spacing_to_zero(footer)

    # Iterate through all tables and their cells
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _set_paragraph_spacing_to_zero(paragraph)
    

    static_dir = os.path.join(os.getcwd(), 'static', 'pdfs')
    
    if not os.path.exists(static_dir):
        os.makedirs(static_dir)

    pdf_file = os.path.join(os.getcwd(), 'static')

    docx_file = os.path.join(static_dir, 'demo2.docx')
    document.save(docx_file)
    pdf_file = os.path.join(static_dir, 'demo')

    #api_key = 'project_public_dd58a2ab023f0c665dc5749a8f0931e0_Pl0dh0a277b8551bb9cdf01e043af64ce0304'

    _convert_to_pdf(docx_file,pdf_file )

    # Return a path relative to the static directory
    return "tradeasy/static/pdfs/demo"

